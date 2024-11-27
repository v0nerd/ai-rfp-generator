import os
from dotenv import load_dotenv
import re
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor

import docx
import pytesseract
from pdf2image import convert_from_bytes
from llama_parse import LlamaParse
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage

load_dotenv()

api_key = os.environ.get("OPENAI_API_KEY")

# def extract_text(file_contents: bytes, file_type: str = "pdf") -> str:
#     if file_type == "pdf":
#         images = convert_from_bytes(file_contents)
#         text = "".join([pytesseract.image_to_string(image) for image in images])
#     elif file_type == "docx":
#         doc = docx.Document(BytesIO(file_contents))
#         text = " ".join([para.text for para in doc.paragraphs])
#     else:
#         text = file_contents.decode("utf-8")
#     return text


def extract_text(file_contents: bytes, file_type: str = "pdf") -> str:
    if file_type == "pdf":
        # Convert PDF to images
        images = convert_from_bytes(file_contents)

        # Use ThreadPoolExecutor to parallelize the OCR task
        with ThreadPoolExecutor() as executor:
            text = "".join(executor.map(pytesseract.image_to_string, images))

    elif file_type == "docx":
        # Process DOCX file
        doc = docx.Document(BytesIO(file_contents))
        text = " ".join([para.text for para in doc.paragraphs])

    else:
        # Assuming it's plain text
        text = file_contents.decode("utf-8")

    return text


def generate_sections(text):
    """
    Generate a technical approach based on given requirements and expertise.

    Parameters:
        requirements (str): A description of the requirements.
        expertise (str): A description of the expertise or domain knowledge.

    Returns:
        str: The generated technical approach content.
    """
    try:
        # Initialize ChatOpenAI with the API key and model configuration
        chat_gpt = ChatOpenAI(
            api_key=api_key,  # Replace with your actual API key or use environment variables for security
            model="gpt-4",
            temperature=0,  # Control creativity (lower = more deterministic)
        )

        # Prepare the message for the AI model
        messages = [
            HumanMessage(
                content=(
                    f"Gnerate sections of the following components in detail from the text:\n"
                    f"Opportunity Details, Project Scope, Requirements, Deliverables\n"
                    f"Text: {text}\n"
                    f"Output a detailed and structured."
                )
            ),
        ]

        # Call the AI model to get a response
        ai_message = chat_gpt.invoke(messages)

        # Display the generated content
        print("Generated Technical Content:", ai_message.content)

        return ai_message.content

    except Exception as e:
        # Handle any errors that occur during the process
        print(f"Error generating technical content: {e}")
        return "An error occurred while generating the technical content."


def extract_text_llama(file_path):
    """
    Use LlamaParser to extract and structure text from the uploaded file.

    Parameters:
        file_path (str): Path to the file to be parsed.

    Returns:
        dict: Parsed sections from the file, including Opportunity Details, Scope, Requirements, and Deliverables.
    """
    try:
        parser = LlamaParse()  # Initialize LlamaParser
        parsed_document = parser.parse(file_path)

        # Extract key sections using LlamaParser's field extraction
        return {
            "Opportunity Details": parsed_document.get_field(
                "Opportunity Details", default="Not Found"
            ),
            "Project Scope": parsed_document.get_field("Scope", default="Not Found"),
            "Requirements": parsed_document.get_field(
                "Requirements", default="Not Found"
            ),
            "Deliverables": parsed_document.get_field(
                "Deliverables", default="Not Found"
            ),
        }
    except Exception as e:
        print(f"Error during LlamaParser processing: {e}")
        return {
            "Opportunity Details": "Not Found",
            "Project Scope": "Not Found",
            "Requirements": "Not Found",
            "Deliverables": "Not Found",
        }


if __name__ == "__main__":
    generate_sections(text)
